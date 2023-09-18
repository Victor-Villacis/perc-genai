import pandas as pd
import tiktoken
from docx import Document  
import openai
from datetime import datetime
import os
import json

def getText(filename):
    # Get all text from word document
    # input: path to word file

    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return "\n".join(fullText)


def split2token(s, n):
    # create a list of sublists
    # input: mainlist,lenght in characters of sublists

    encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
    encoded_text = encoding.encode(s)
    text_list = [
        encoding.decode(encoded_text[i : i + n]) for i in range(0, len(encoded_text), n)
    ]
    return text_list


### Add this function by Victor Wang to resolve the probelm
# def splitsections():


def text_to_frame(textList):
    # Convert list into data frame and parse out only the big sections
    # This part can be refinede depending on process
    df_text = pd.DataFrame(columns=["section", "section_text"])
    for para in range(1, len(textList)):
        if len(textList[para]) > 100:
            if (
                not str(textList[para - 1].split("//]")[0])
                in df_text["section"].unique()
            ):
                # print(str(textList[para-1].split("//]")[0]))
                temp_df = pd.DataFrame()
                temp_df["section"] = [str(textList[para - 1].split("//]")[0])]
                temp_df["section_text"] = [textList[para]]
                df_text = pd.concat([df_text, temp_df])
    return df_text


def map2text(outlist, titles=None):
    import collections, ast, textwrap

    titles = ["Disease", "Experience/opinion/knowledge", "Quote 1", "Quote 2", "others"]
    # outmap = collections.defaultdict(list)
    outstr = ""
    for x, y in outlist.items():
        if x != "NA" and x is not None:
            outstr = outstr + textwrap.indent(x, " " * 5) + "\n"
            if isinstance(y, dict):
                for k, v in y.items():
                    if isinstance(v, str):
                        outstr = outstr + textwrap.indent(k + ": ", " " * 10) + v + "\n"
                    elif v != "NA" and isinstance(v, list):
                        outstr = outstr + textwrap.indent("Name: ", " " * 10) + k + "\n"
                        for i, element in enumerate(v):
                            outstr = (
                                outstr
                                + textwrap.indent(titles[i] + ": ", " " * 15)
                                + element
                                + "\n"
                            )
    return outstr


def run_promt(req, text_inp, engine, max_tokens_num=2000):
    # Run ChatGPT promt
    # input: text to be used inside promt to work as a question for ChatGPT, text to be summaried, the different promts
    # output the second level promts outputs
    import openai

    response = openai.ChatCompletion.create(
        engine=engine,  # "gpt-35-turbo-16k"
        messages=[
            {
                "role": "system",
                "content": """You are an AI assistant with expertise in patient engagement specialist for a pharmaceutical. 
            You will be given a body of text consisting patient engagement converations""",
            },
            {
                "role": "user",
                "content": """
            --
            %s
            --
            %s"""
                % (text_inp, req),
            },
        ],
        temperature=0.2,
        max_tokens=max_tokens_num,
        top_p=0.1,
        frequency_penalty=0,
        presence_penalty=0,
        stop=None,
    )
    print(response["usage"])

    return response["choices"][0]["message"]["content"]


def summarize_file(filename, timestamp, folder, P_1_L_1, P_1_L_2):
    # main function to summarize all text for a given file
    # input: file path to be summarized, current timestamp, folder for output
    import re, textwrap, openai, time, ast
    from collections import defaultdict

    text = getText(filename)

    text = text.replace("\n", "").replace("'", "").replace("?", "? ")
    text_list = re.split(r"\[//.*?//\]", text)
    sections = re.findall(r"\[//.*?//\]", text)
    sections = [x.replace("[//", "").replace("//]", "") for x in sections]
    text_list = [x for x in text_list if len(x) > 200]
    # df = text_to_frame(text_list).reset_index(drop=True)
    openai.api_base = "https://genaiapimna.jnj.com/openai-chat"
    openai.api_version = (
        "2023-03-15-preview"  # #  only "2023-05-15" works for encrypted subs
    )
    openai.api_key = "228b7c3ed183460abb331208fd2893b3"
    engine1 = "gpt-35-turbo-16k"

    sub_list = [split2token(x, 13500) for x in text_list]

    summary_list = defaultdict(list)
    # summary_list2 = []
    for p in range(len(sub_list)):  # len(sub_list)
        for j in range(len(sub_list[p])):  # len(sub_list[p])
            sub_file = sub_list[p][j]
            print("The segment {j} in section {p}".format(j=j + 1, p=p + 1))
            try:
                # First promts
                resp = run_promt(P_1_L_1, sub_file, engine1)
                summary_list[sections[p]].append(resp)
                time.sleep(1)
            except openai.error.InvalidRequestError as e:
                # Handle connection error here
                print(f"Failed because of invalid requests: {e}")
                pass

            except:
                # if first attempts fails - wait and make one retry
                time.sleep(10)
                resp = run_promt(P_1_L_1, sub_file, engine1)
                summary_list[sections[p]].append(resp)
                time.sleep(2)

        # print(resp)
    # outputs level 1 summary incase needed

    level1sum = ""
    for k, v in summary_list.items():
        level1sum = level1sum + k + "\n"
        for s in v:
            s = ast.literal_eval(s)
            if isinstance(s, dict):
                level1sum = level1sum + map2text(s) + "\n"

            else:
                level1sum = "".join(summary_list)
    # with open(
    #     outputfolder + "/The interview summary L1 - " + dt_string + ".txt", "w+"
    # ) as f:
    #     f.write(level1sum)

    # with open(outputfolder + "/The first level of Summary20230812121945.txt", 'r+') as f:
    #     level1sum = f.readlines()

    # summary = []
    # sdoc = Document()
    # time.sleep(30)
    openai.api_base = "https://ims-openai-qa.openai.azure.com/"
    openai.api_version = "2023-06-01-preview"
    openai.api_key = "e0a02ca2e58041109e2c4c27e19a4a87"
    engine2 = "gpt-4-32k"
    sec_summary = run_promt(P_1_L_2, level1sum, engine=engine2, max_tokens_num=5000)
    # table = pd.read_json(json.loads(sec_summary))
    level2sum = map2text(ast.literal_eval(sec_summary))
    # sdoc.add_paragraph(level2sum)

    # with open(
    #     outputfolder + "/The interview summary L2 - " + dt_string + ".txt", "w+"
    # ) as f:
    #     f.write(level2sum)

    # sdoc.save(outputfolder + "/The second level of Summary" +dt_string + ".docx")
    """ 
    try:
        summary_map != NA
        top_summary = run_promt(P_1_L_2,''.join(summary))
    except:
        top_summary = ''.join(summary) 
    """
    return summary_list, sec_summary

    ### To save time and cost, leverage the first level of summary by running the following code
    # text = getText(outputfolder + "/The second level of Summary20230719180322.docx")
    # top_summary = run_promt(P_1_L_2,text)
    # return top_summary


# Promts being used for general summary purpose

# %%


def main():
  
    script_dir = os.path.dirname(os.path.realpath(__file__))

    inputfolder = os.path.join(script_dir, "../..", "server/src/services", "input_folder")
    outputfolder = os.path.join(script_dir, "../..", "server/src/services", "output_folder")

    if not os.path.exists(inputfolder):
        os.makedirs(inputfolder)
    if not os.path.exists(outputfolder):
        os.makedirs(outputfolder)


    hosts = ["Wes P", "Gabrielle G", "Bridget D"]
    hs = ", ".join(hosts)

    now = datetime.now()
    dt_string = now.strftime("%Y%m%d%H%M%S")

    files = [f for f in os.listdir(inputfolder) if os.path.isfile(os.path.join(inputfolder, f))]
    if len(files) == 0:
        print("No files found in the input folder.")
        return

    # Pick the first file
    single_file = files[0]
    # Construct the full path
    file_path = os.path.join(inputfolder, single_file)

    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        file_content = f.read()


    # constants
    openai.api_type = "azure"

    prompt_1_level_1 = """ Summarize the conversation discussed here into different topics with very high details. For each topic, summary the topic with high details. List two quotes from participants for each topic (exlude any quote from these participants: %s). For each topic, structure your output using the following format as an example. Please return your output in Json:
            {
            "<The first topic>": {
                "Summary": "<summarize this topic with very high details>",  
                "Quote 1 - <first name, the initial of the last name>": "<quote one sentence from a participant>", 
                "Quote 2 - <first name, the initial of the last name>": "<quote another sentence from a second participant"},
            "<The second topic>": { 
                "Summary": "<summarize this topic with very high details>",  
                "Quote 1 - <first name, the initial of the last name>": "<quote one sentence from a participant", 
                "Quote 2 - <first name, the initial of the last name>": "<quote another sentence from a second participant"}
            
            }
    """ % (
        hs
    )
    prompt_1_level_2 = """Your input has the structure as sections, topics, topic summaries, 2 quotes from the participants about each topic. Please group all summaries in the text into 5-8 different topics in total by pooling similar ones together. Then for each topic, summaryize the topic and cite 2 quotes. Please return your output in Json.

            
            {
            "<The first topic>": {
                "Summary": "<summarize this topic with very high details>",  
                "Quote 1 - <first name, the initial of the last name>": "<quote one sentence from a participant>", 
                "Quote 2 - <first name, the initial of the last name>": "<quote another sentence from a second participant"},
            "<The second topic>": { 
                "Summary": "<summarize this topic with very high details>",  
                "Quote 1 - <first name, the initial of the last name>": "<quote one sentence from a participant", 
                "Quote 2 - <first name, the initial of the last name>": "<quote another sentence from a second participant"}
            
            }



    """
    level1, level2 = summarize_file(
        file_path, dt_string, outputfolder, prompt_1_level_1, prompt_1_level_2
    )


    print(level1)
    print("********************")
    print(level2)

    # Save level1 and level2 as JSON files in the output folder
    level1_path = os.path.join(outputfolder, "level_1.json")
    level2_path = os.path.join(outputfolder, "level_2.json")

    with open(level1_path, "w", encoding='utf-8') as f1:
        json.dump(level1, f1, ensure_ascii=False, indent=4)

    with open(level2_path, "w", encoding='utf-8') as f2:
        json.dump(level2, f2, ensure_ascii=False, indent=4)

    return level1, level2


if __name__ == "__main__":
    main()